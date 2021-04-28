from selenium import webdriver
from time import sleep
from docx import Document
from transliterate import translit


def sort_by_length(input_string):
    return len(input_string)


def get_a_document_with_an_empty_table_to_fill_in(elibrary_page):
    # кол-во строк в таблице = наибольшему кол-ву учёных на одном из сайтов (на elibrary.ru их больше, берём оттуда)
    xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]' \
            '/table/tbody/tr/td/div[3]/table/tbody/tr/td[2]/font/b[1]/font'
    # добавляем к кол-ву строк в таблице единицу, т.к. первая строка - поля таблицы
    number_of_rows_in_the_table = int(elibrary_page.find_element_by_xpath(xpath=xpath).text)

    document_with_a_table = Document()
    document_with_a_table.add_paragraph('Работу выполнил студент гр. ПИбд-11, Курмыза Павел Сергеевич.')
    document_with_a_table.add_paragraph('Сравнительный анализ сайтов elibrary.ru и scholar.google.ru по цитированиям '
                                        'учёных УлГТУ', 'Title')
    document_with_a_table.add_paragraph('Таблица 1. Список всех учёных с обоих сайтов:')

    # добавляем таблицу размерностью (number_of_rows_in_the_table + 1)x5 (+1, т.к. первая строка - поля таблицы)
    table = document_with_a_table.add_table(rows=number_of_rows_in_the_table + 1, cols=5)

    # применяем стиль для таблицы
    table.style = 'Table Grid'

    first_column = table.cell(0, 0)
    first_column.text = 'Порядковый номер'

    second_column = table.cell(0, 1)
    second_column.text = 'Учёный (eLIBRARY.ru)'

    third_column = table.cell(0, 2)
    third_column.text = 'Количество цитирований (eLIBRARY.ru)'

    second_column = table.cell(0, 3)
    second_column.text = 'Учёный (Академия Google)'

    third_column = table.cell(0, 4)
    third_column.text = 'Количество цитирований (Академия Google)'

    # заполняем таблицу данными
    for i in range(4, number_of_rows_in_the_table + 4):
        # получаем (i-3) ячейку 1-го (0-го) столбца
        cell = table.cell(i - 3, 0)
        # записываем в 1-й (0-й) столбец номера научных деятелей
        cell.text = str(i - 3)

    # возвращаем table, чтобы таблицу наполнить данными, document_with_a_table, чтобы сохранить наполненную таблицу
    return document_with_a_table, table


def log_in_elibrary_and_go_to_the_scientists():
    option = webdriver.FirefoxOptions()
    option.set_preference('dom.webdriver.enabled', False)
    option.set_preference('general.useragent.override', 'Pablo')

    browser = webdriver.Firefox(options=option)

    browser.get('https://www.elibrary.ru/defaultx.asp')
    # вводим логин
    entering_login_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[1]/table' \
                           '/tbody/tr[4]/td/div/div/table[1]/tbody/tr[6]/td/div/input'
    browser.find_element_by_xpath(xpath=entering_login_xpath).click()
    browser.find_element_by_xpath(xpath=entering_login_xpath).send_keys('tmrrwnxtsn')
    sleep(3)

    # вводим пароль
    entering_pass_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[1]/table' \
                          '/tbody/tr[4]/td/div/div/table[1]/tbody/tr[8]/td/div/input'
    browser.find_element_by_xpath(xpath=entering_pass_xpath).click()
    browser.find_element_by_xpath(xpath=entering_pass_xpath).send_keys('78ofWCM=Pogp')
    sleep(3)

    # нажимаем "войти"
    entering_btn_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[1]/table' \
                         '/tbody/tr[4]/td/div/div/table[1]/tbody/tr[9]/td/div[2]'
    browser.find_element_by_xpath(xpath=entering_btn_xpath).click()
    sleep(3)

    # переходим в раздел "авторы"
    scientist_btn_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[1]/table' \
                        '/tbody/tr[4]/td/div/div/table[1]/tbody/tr[5]/td[2]/a'
    browser.find_element_by_xpath(xpath=scientist_btn_xpath).click()
    sleep(3)

    # возвращаем страницу сайта, из которой мы будем брать информацию
    return browser


def fill_in_the_table_with_data_from_elibrary(elibrary_page, document_with_a_table, table):
    next_page_controller = 0

    # количество учёных на elibrary.ru
    scientists_count = len(table.rows) - 1

    for j in range(3, ((scientists_count // 100) + 4)):

        # переходим на следующую страницу с учёными
        next_page_btn_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]/table' \
                              f'/tbody/tr/td/div[4]/table/tbody/tr/td[{j}]'
        elibrary_page.find_element_by_xpath(xpath=next_page_btn_xpath).click()
        sleep(3)

        # смотрим --> "Показано на данной странице: с start_scientist_number по end_scientist_number."
        start_scientist_number_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]' \
                                       '/table/tbody/tr/td/div[3]/table/tbody/tr/td[2]/font/b[3]'
        start_scientist_number = int(elibrary_page.find_element_by_xpath(xpath=start_scientist_number_xpath).text)
        end_scientist_number_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]' \
                                     '/table/tbody/tr/td/div[3]/table/tbody/tr/td[2]/font/b[4]'
        end_scientist_number = int(elibrary_page.find_element_by_xpath(xpath=end_scientist_number_xpath).text)

        scientist_count_on_page = end_scientist_number - start_scientist_number + 1

        # в данном разделе извлекаем следующую информацию:
        for i in range(4, scientist_count_on_page + 4):
            scientist_name_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table' \
                                   f'/tbody/tr[2]/td[1]/table/tbody/tr/td/table/tbody/tr[{i}]/td[3]'
            scientist_name = elibrary_page.find_element_by_xpath(xpath=scientist_name_xpath)
            cell = table.cell(i - 3 + next_page_controller, 1)
            cell.text = scientist_name.text.replace('Ульяновский государственный'
                                                    ' технический университет '
                                                    '(Ульяновск)', '').replace('*', '').replace('  ', ' ').strip()

            number_of_citations_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table' \
                                        f'/tbody/tr[2]/td[1]/table/tbody/tr/td/table/tbody/tr[{i}]/td[5]'
            number_of_citations = elibrary_page.find_element_by_xpath(xpath=number_of_citations_xpath)
            cell = table.cell(i - 3 + next_page_controller, 2)
            cell.text = number_of_citations.text

        next_page_controller += 100

    document_with_a_table.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    elibrary_page.quit()


def get_information_from_elibrary():
    elibrary_page = log_in_elibrary_and_go_to_the_scientists()
    document_with_a_table, table = get_a_document_with_an_empty_table_to_fill_in(elibrary_page)
    fill_in_the_table_with_data_from_elibrary(elibrary_page, document_with_a_table, table)


def go_to_google_scholar():
    option = webdriver.FirefoxOptions()
    option.set_preference('dom.webdriver.enabled', False)
    option.set_preference('general.useragent.override', 'Pablo')

    browser = webdriver.Firefox(options=option)
    browser.get('https://scholar.google.ru/citations?hl=ru&view_op=search_authors&mauthors=улгту&btnG=')

    return browser


def get_a_document_with_a_table_to_fill_in():
    doc = Document('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    table = doc.tables[0]
    return doc, table


def fill_in_the_table_with_data_from_scholar(scholar_page, document_with_a_table, table):
    next_page_controller = 0

    while True:
        first_and_last_scientist_numbers = scholar_page.find_element_by_class_name('gsc_pgn_ppn').text
        first_scientist_number_on_page, last_scientist_number_on_page = map(int,
                                                                            first_and_last_scientist_numbers.split(
                                                                                ' - '))

        for i in range(1, (last_scientist_number_on_page - first_scientist_number_on_page + 2)):
            scientist_name_xpath = f'/html/body/div/div[8]/div[2]/div/div[{i}]/div/div/h3/a'
            scientist_name = scholar_page.find_element_by_xpath(xpath=scientist_name_xpath)
            cell = table.cell(i + next_page_controller, 3)
            cell.text = scientist_name.text

            number_of_scientist_citations_xpath = f'/html/body/div/div[8]/div[2]/div/div[{i}]/div/div/div[3]'
            number_of_scientist_citations = scholar_page.find_element_by_xpath(
                xpath=number_of_scientist_citations_xpath)
            cell = table.cell(i + next_page_controller, 4)

            if number_of_scientist_citations.text == '':
                cell.text = '0'
            else:
                cell.text = number_of_scientist_citations.text.replace('Цитируется: ', '')

        try:  # пытаемся перейти на следующую страницу с учёными
            next_page_btn_xpath = '/html/body/div/div[8]/div[2]/div/div[11]/div/button[2]'
            scholar_page.find_element_by_xpath(xpath=next_page_btn_xpath).click()
            next_page_controller += 10
            sleep(3)
        except Exception as e:  # если не удалось перейти на следующую страницу, то завершаем работу со scholar
            print(e)
            break

    document_with_a_table.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    scholar_page.quit()


def get_information_from_google_scholar():
    scholar_page = go_to_google_scholar()
    document_with_a_table, table = get_a_document_with_a_table_to_fill_in()
    fill_in_the_table_with_data_from_scholar(scholar_page, document_with_a_table, table)


def get_the_result_of_the_analysis_of_indicators(elibrary_indicator, scholar_indicator):
    result = ''
    counter = 0
    # сначала определим, насколько существенна разница между показателями
    delta = abs(elibrary_indicator - scholar_indicator)
    if delta == 0:
        result += 'Показатели равны. На обоих сайтах одинаковое количество цитирований учёного.'
    else:
        if elibrary_indicator > scholar_indicator:
            counter += 1
            result += f' На elibrary.ru цитирований учёного больше на {delta}'
            if scholar_indicator != 0 and elibrary_indicator / scholar_indicator > 1.5:
                result += f' (Примерно в {round((elibrary_indicator / scholar_indicator), 2)} р.).'
            else:
                result += '.'
        else:
            counter -= 1
            result += f' На scholar.google.ru цитирований учёного больше на {delta}'
            if elibrary_indicator != 0 and scholar_indicator / elibrary_indicator > 1.5:
                result += f' (Примерно в {round((scholar_indicator / elibrary_indicator), 2)} р.).'
            else:
                result += '.'
    return result, counter


def determine_which_site_has_more_information(document, number):
    # определяет, на каком сайте больше информации про учёных если > 0, то на elibrary, если < 0, то на scholar,
    # если = 0, то одинаково и там, и там
    document.add_paragraph('')
    if number > 0:
        document.add_paragraph('На основе сравнительного анализа 2-ой таблицы можно сделать вывод: сайт elibrary.ru'
                               ' предоставляет более обширную информацию про учёных УлГТУ.')
    elif number < 0:
        document.add_paragraph('На основе сравнительного анализа 2-ой таблицы можно сделать вывод: сайт scholar.google.'
                               'ru предоставляет более обширную информацию про учёных УлГТУ.')
    else:
        document.add_paragraph('На основе сравнительного анализа 2-ой таблицы можно сделать вывод: сайты elibrary.ru и '
                               ' scholar.google.ru предоставляют информацию про учёных УлГТУ примерно одинакового '
                               'объёма.')


def create_the_second_table(document):
    # добавляем таблицу размерностью (number_of_rows_in_the_table + 1)x5 (+1, т.к. первая строка - поля таблицы)
    document.add_paragraph('Таблица 2. Учёные, которые есть на обоих сайтах, с указанием цитирований:')

    table = document.add_table(rows=1, cols=6)

    # применяем стиль для таблицы
    table.style = 'Table Grid'

    first_column = table.cell(0, 0)
    first_column.text = '№'

    second_column = table.cell(0, 1)
    second_column.text = 'Учёный (Как записан на elibrary.ru)'

    third_column = table.cell(0, 2)
    third_column.text = 'Количество цитирований (elibrary.ru)'

    fourth_column = table.cell(0, 3)
    fourth_column.text = 'Учёный (Как записан на scholar.google.ru)'

    fifth_column = table.cell(0, 4)
    fifth_column.text = 'Количество цитирований (scholar.google.ru)'

    sixth_column = table.cell(0, 5)
    sixth_column.text = 'Результат сравнения показателей'

    return table


def process_first_table_and_create_second_table_based_on_it():
    doc = Document('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    first_table = doc.tables[0]

    elibrary_scientists_array = []
    elibrary_scientists_citations_array = []

    sum_elibrary_citations = 0
    sum_elibrary_scientists = 0

    sum_scholar_citations = 0
    count_of_scholar_scientists = 0

    for row in first_table.rows[1:]:
        # при сопоставлении показателей одних и тех же учёных на разных сайтах мы будем работать только с теми,
        # у кого и на elibrary, и на scholar цитирования есть и кол-во цитирований больше 0
        if row.cells[1].text != '' and row.cells[2].text != '0':
            elibrary_scientist_citations_number = int(row.cells[2].text)
            sum_elibrary_citations += elibrary_scientist_citations_number

            elibrary_scientist = row.cells[1].text
            elibrary_scientists_array.append(elibrary_scientist)
            elibrary_scientists_citations_array.append(elibrary_scientist_citations_number)

        # считаем всех учёных, даже тех, у кого нет цитирований, поэтому из верхнего условия выносим
        sum_elibrary_scientists += 1

        # считаем общее количество авторов scholar
        if row.cells[3].text != '':  # если в ячейке есть имя автора, т.е. её значени не равно "пустой строке", то
            count_of_scholar_scientists += 1

            scholar_scientist_citations_number = int(row.cells[4].text)
            sum_scholar_citations += scholar_scientist_citations_number

    doc.add_paragraph('\n')
    doc.add_paragraph(
        f'Суммарное количество цитирований учёных УлГТУ (по данным elibrary.ru): {sum_elibrary_citations}.')
    doc.add_paragraph(
        f'Суммарное количество цитирований учёных УлГТУ (по данным scholar.google.ru): {sum_scholar_citations}.')

    if sum_elibrary_citations > sum_scholar_citations:
        doc.add_paragraph(f'В базе elibrary.ru цитирований учёных УлГТУ больше, чем в базе scholar.google.ru, на '
                          f'{sum_elibrary_citations - sum_scholar_citations} (Примерно в '
                          f'{round(sum_elibrary_citations / sum_scholar_citations, 2)} р. больше).')
    elif sum_scholar_citations > sum_elibrary_citations:
        doc.add_paragraph(f'В базе scholar.google.ru цитирований учёных УлГТУ больше, чем в базе elibrary.ru, на '
                          f'{sum_scholar_citations - sum_elibrary_citations} (Примерно в '
                          f'{round(sum_scholar_citations / sum_elibrary_citations, 2)} р. больше).')
    else:
        doc.add_paragraph('В базах elibrary.ru и scholar.google.ru одинаковое количество цитирований учёных УлГТУ.')

    doc.add_paragraph(f'Суммарное количество учёных УлГТУ (по данным elibrary.ru): {sum_elibrary_scientists}.')
    doc.add_paragraph(
        f'Суммарное количество учёных УлГТУ (по данным scholar.google.ru): {count_of_scholar_scientists}.')

    if sum_elibrary_scientists > count_of_scholar_scientists:
        doc.add_paragraph(f'В базе elibrary.ru учёных УлГТУ больше, чем в базе scholar.google.ru, на '
                          f'{sum_elibrary_scientists - count_of_scholar_scientists} (Примерно в '
                          f'{round(sum_elibrary_scientists / count_of_scholar_scientists, 2)} р. больше).')
    elif sum_scholar_citations > sum_elibrary_citations:
        doc.add_paragraph(f'В базе scholar.google.ru учёных УлГТУ больше, чем в базе elibrary.ru, на '
                          f'{count_of_scholar_scientists - sum_elibrary_scientists} (Примерно в '
                          f'{round(count_of_scholar_scientists / sum_elibrary_scientists, 2)} р. больше).')
    else:
        doc.add_paragraph('В базах elibrary.ru и scholar.google.ru одинаковое количество учёных УлГТУ.')

    # создаём вторую таблицу, в которую мы поместим "совпавших" на обоих сайтах учёных, их показатель и результат
    # сравнения
    second_table = create_the_second_table(doc)

    count_of_scientists = 0

    # переменная для отслеживания: какой сайт обладает более обширной информацией
    which_site_has_the_most_information = 0
    for i in range(1, count_of_scholar_scientists + 1):
        scholar_scientist = first_table.rows[i].cells[3].text
        scholar_scientist_citations = first_table.rows[i].cells[4].text

        if scholar_scientist[0].lower() in 'qwertyuiopasdfghjklzxcvbnm':
            # перевод имён, написанных по-русски латиницей, в кириллицу
            scholar_scientist = translit(scholar_scientist, 'ru')

        scholar_scientist = sorted(scholar_scientist.split(), key=sort_by_length)
        # ищем учёных, которые есть на обоих сайтах (проходимся по всем учёным elibrary, т.к. на elb их
        # значительно больше)
        for j in range(len(elibrary_scientists_array)):
            elibrary_scientist = sorted(elibrary_scientists_array[j].split(), key=sort_by_length)
            # если вхождения 2-х или более слов имени учёного на scholar есть в имени учёного на elibrary
            # если хотя бы 2 слова из имени учёного на scholar 1 в 1 совпадают с 2-мя словами имени учёного на elb
            if set(scholar_scientist).issubset(elibrary_scientist) or set(elibrary_scientist).issubset(
                    scholar_scientist):
                # считаем учёных, которые есть и на elibrary, и на scholar.google
                count_of_scientists += 1

                # добавляем строку в таблицу (в конец)
                second_table.add_row()

                # записываем в 1 столбец порядковый номер учёного в таблице
                cell = second_table.cell(count_of_scientists, 0)
                cell.text = str(count_of_scientists)

                # записываем во 2 столбец ФИО учёного (как записан на elibrary)
                cell = second_table.cell(count_of_scientists, 1)
                cell.text = elibrary_scientists_array[j]

                # записываем в 3 столбец цитирования учёного по данным elibrary
                cell = second_table.cell(count_of_scientists, 2)
                cell.text = str(elibrary_scientists_citations_array[j])

                # записываем в 4 столбец ФИО учёного (как записан на scholar)
                cell = second_table.cell(count_of_scientists, 3)
                cell.text = ' '.join(scholar_scientist)

                # записываем в 5 столбец цитирования учёного по данным scholar
                cell = second_table.cell(count_of_scientists, 4)
                cell.text = str(scholar_scientist_citations)

                # записываем в 6 столбец результат сравнительного анализа
                cell = second_table.cell(count_of_scientists, 5)
                cell.text, counter_which_site = get_the_result_of_the_analysis_of_indicators(
                    int(elibrary_scientists_citations_array[j]),
                    int(scholar_scientist_citations))

                which_site_has_the_most_information += counter_which_site
                # дальше не ищем, и переходим к новому учёному
                break

    # дописываем в конец ворд-документа заключительный вывод сравнительного анализа данных
    determine_which_site_has_more_information(doc, which_site_has_the_most_information)

    doc.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')


def get_more_extensive_elibrary_scientists_information():
    elibrary_dictionary = dict()

    # "логинимся" на elibrary.ru и переходим в раздел "авторы"
    browser = log_in_elibrary_and_go_to_the_scientists()

    doc = Document('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    first_table = doc.tables[0]

    # учёные elibrary (столбец с индексом 1 - row.cells[1]), чьи профили есть на 2 сайтах
    second_table = doc.tables[1]
    array_of_scientists_second_table = []
    for row in second_table.rows[1:]:
        array_of_scientists_second_table.append(row.cells[1].text)

    doc.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')

    next_page_controller = 0

    # количество учёных elibrary.ru
    scientists_count = len(first_table.rows) - 1

    for j in range(3, ((scientists_count // 100) + 4)):

        # переходим на следующую страницу с учёными
        next_page_btn_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]/table' \
                              f'/tbody/tr/td/div[4]/table/tbody/tr/td[{j}]'
        browser.find_element_by_xpath(xpath=next_page_btn_xpath).click()
        sleep(3)

        # смотрим --> "Показано на данной странице: с start_scientist_number по end_scientist_number."
        start_scientists_number_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]' \
                                        '/table/tbody/tr/td/div[3]/table/tbody/tr/td[2]/font/b[3]'
        start_scientist_number = int(browser.find_element_by_xpath(xpath=start_scientists_number_xpath).text)
        end_scientist_number_xpath = '/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]/td[1]' \
                                     '/table/tbody/tr/td/div[3]/table/tbody/tr/td[2]/font/b[4]'
        end_scientist_number = int(browser.find_element_by_xpath(xpath=end_scientist_number_xpath).text)

        scientists_count_on_page = end_scientist_number - start_scientist_number + 1

        # в данном разделе извлекаем следующую информацию:
        for i in range(4, scientists_count_on_page + 4):
            # получаем ФИО учёного
            scientist_name_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table' \
                                   f'/tbody/tr[2]/td[1]/table/tbody/tr/td/table/tbody/tr[{i}]/td[3]'
            scientist_name = browser.find_element_by_xpath(xpath=scientist_name_xpath).text.replace(
                'Ульяновский государственный'
                ' технический университет '
                '(Ульяновск)', '').replace('*', '').replace('  ', ' ').strip()
            # если учёный есть среди тех, чьи профили присутствуют на обоих сайтах
            if scientist_name in array_of_scientists_second_table:
                # получаем количество его публикаций
                publications_count_xpath = f'/html/body/table/tbody/tr/td/table[1]/tbody/tr/td[2]/table/tbody/tr[2]' \
                                           f'/td[1]/table/tbody/tr/td/table/tbody/tr[{i}]/td[4]/div/a[1]'
                publications_count = int(browser.find_element_by_xpath(xpath=publications_count_xpath).text)
                if publications_count >= 5:
                    # переходим к его публикациям
                    browser.find_element_by_xpath(xpath=publications_count_xpath).click()
                    sleep(15)

                    article_names = list()
                    # проходимся по первым 5 публикациям (самым цитируемым)
                    for k in range(4, 9):
                        flag = 1
                        while flag:
                            try:
                                article_name_xpath = f'/html/body/div[3]/table/tbody/tr/td/table[1]/tbody/tr/td[2]' \
                                                     f'/form/table/tbody/tr[2]/td[1]/table/tbody/tr/td/table/tbody' \
                                                     f'/tr[{k}]/td[2]'
                                article_name = browser.find_element_by_xpath(xpath=article_name_xpath).text.split('\n')[
                                    0].title()

                                article_names.append(article_name)

                                flag = 0
                            except Exception as e:
                                print(e)
                                browser.refresh()
                                sleep(10)

                    elibrary_dictionary[scientist_name] = article_names
                    browser.back()
                    sleep(3)

        next_page_controller += 100

    browser.quit()

    return elibrary_dictionary


def get_more_extensive_scholar_scientists_information():
    scholar_dictionary = dict()

    option = webdriver.FirefoxOptions()
    option.set_preference('dom.webdriver.enabled', False)
    option.set_preference('general.useragent.override', 'Pablo')

    browser = webdriver.Firefox(options=option)
    browser.get('https://scholar.google.ru/citations?hl=ru&view_op=search_authors&mauthors=улгту&btnG=')

    doc = Document('Сравнение показателей elibrary.ru and scholar.google.ru.docx')

    # учёные scholar.google.ru (столбец с индексом 3 - row.cells[3]), чьи профили есть на 2 сайтах
    second_table = doc.tables[1]
    array_of_scientists_second_table = []
    for row in second_table.rows[1:]:
        array_of_scientists_second_table.append(row.cells[3].text)

    doc.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')

    next_page_controller = 0

    while True:
        first_and_last_scientists_numbers = browser.find_element_by_class_name('gsc_pgn_ppn').text
        first_scientist_number_on_page, last_scientist_number_on_page = map(int,
                                                                            first_and_last_scientists_numbers.split(
                                                                                ' - '))

        for i in range(1, (last_scientist_number_on_page - first_scientist_number_on_page + 2)):
            scientist_name_xpath = f'/html/body/div/div[8]/div[2]/div/div[{i}]/div/div/h3/a'
            scientist_name = browser.find_element_by_xpath(xpath=scientist_name_xpath).text

            if scientist_name[0].lower() in 'qwertyuiopasdfghjklzxcvbnm':
                # перевод имён, написанных по-русски латиницей, в кириллицу
                scientist_name = translit(scientist_name, 'ru')

            # так как во 2 таблице расположенны отсортированные по длине имена учёных, то, чтобы сравнивать имена,
            # полученные с сайта, нужно их тоже отсортировать аналогичным способом, а потом только сравнивать
            scientist_name = sorted(scientist_name.split(), key=sort_by_length)
            scientist_name = ' '.join(scientist_name)

            if scientist_name in array_of_scientists_second_table:
                browser.find_element_by_xpath(xpath=scientist_name_xpath).click()
                sleep(10)

                articles_count = 0
                article_names = list()
                # проходимся по первым 5 публикациям (самым цитируемым) 4-мя разными способами, т.к.
                # структура страниц учёных отличается
                for k in range(1, 6):
                    flag = 1
                    while flag:
                        try:
                            article_name_xpath = f'/html/body/div/div[13]/div[2]/div/div[4]/form/div[1]/table/tbody' \
                                                 f'/tr[{k}]/td[1]/a'
                            article_name = browser.find_element_by_xpath(xpath=article_name_xpath).text.title()

                            if article_name not in article_names:
                                article_names.append(article_name)
                                articles_count += 1
                            flag = 0
                        except Exception as e:
                            print(e)
                            sleep(2)
                            try:
                                article_name_xpath = f'/html/body/div/div[13]/div[2]/div/div[3]/form/div[1]/table/' \
                                                     f'tbody/tr[{k}]/td[1]/a'
                                article_name = browser.find_element_by_xpath(xpath=article_name_xpath).text.title()

                                if article_name not in article_names:
                                    article_names.append(article_name)
                                    articles_count += 1
                                flag = 0
                            except Exception as e:
                                print(e)
                                sleep(2)
                                try:
                                    article_name_xpath = f'/html/body/div/div[13]/div[2]/div/div[3]/form/div[1]/table' \
                                                         f'/tbody/tr/td[1]/a'
                                    article_name = browser.find_element_by_xpath(xpath=article_name_xpath).text.title()

                                    if article_name not in article_names:
                                        article_names.append(article_name)
                                        articles_count += 1
                                    flag = 0
                                except Exception as e:
                                    print(e)
                                    sleep(2)
                                    try:
                                        article_name_xpath = f'/html/body/div/div[13]/div[2]/div/div[4]/form/div[1]' \
                                                             f'/table/tbody/tr/td[1]/a'
                                        article_name = browser.find_element_by_xpath(
                                            xpath=article_name_xpath).text.title()

                                        if article_name not in article_names:
                                            article_names.append(article_name)
                                            articles_count += 1
                                        flag = 0
                                    except Exception as e:
                                        print(e)
                                        sleep(2)
                                        break
                if articles_count == 5:
                    scholar_dictionary[scientist_name] = article_names
                browser.back()
                sleep(4)

        try:  # пытаемся перейти на следующую страницу с авторами
            next_page_btn_xpath = '/html/body/div/div[8]/div[2]/div/div[11]/div/button[2]'
            browser.find_element_by_xpath(xpath=next_page_btn_xpath).click()
            next_page_controller += 10
            sleep(3)
        except Exception as e:  # если не удалось перейти на следующую страницу, то завершаем работу со scholar
            print(e)
            break
    browser.quit()

    return scholar_dictionary


def compare_dictionaries(elibrary_dictionary, scholar_dictionary):
    groups = {'Самые цитируемые публикации полностью совпадают': [],
                'Самые цитируемые публикации почти совпадают': [],
                'Самые цитируемые публикации почти не совпадают': [],
                'Самые цитируемые публикации полностью не совпадают': []}
    for scholar_scientist, scholar_scientist_articles in sorted(scholar_dictionary.items()):

        array_words_in_scholar_scientist = scholar_scientist.split()

        for elibrary_scientist, elibrary_scientist_articles in sorted(elibrary_dictionary.items()):

            # для удобного сравнения имён расположим слова в имени учёного на elibrary в порядке возрастания длины, т.к.
            # в этом же порядке расположены слова в имени учёного на scholar
            array_words_in_elibrary_scientist = sorted(elibrary_scientist.split(), key=sort_by_length)

            # входят ли все слова из имени на scholar в имя на elbr или все слова из имени на elbrв имя на scholar
            if set(array_words_in_scholar_scientist).issubset(array_words_in_elibrary_scientist) or \
                    set(array_words_in_elibrary_scientist).issubset(array_words_in_scholar_scientist):
                if len(elibrary_dictionary[elibrary_scientist]) == 5 and len(
                        scholar_dictionary[scholar_scientist]) == 5:
                    number_of_matches = 0
                    for i in list(scholar_scientist_articles):
                        for j in list(elibrary_scientist_articles):
                            if i == j:
                                number_of_matches += 1

                    if number_of_matches == 5:
                        groups['Самые цитируемые публикации полностью совпадают'].append(elibrary_scientist)
                    elif number_of_matches == 3 or number_of_matches == 4:
                        groups['Самые цитируемые публикации почти совпадают'].append(elibrary_scientist)
                    elif number_of_matches == 1 or number_of_matches == 2:
                        groups['Самые цитируемые публикации почти не совпадают'].append(elibrary_scientist)
                    else:
                        groups['Самые цитируемые публикации полностью не совпадают'].append(elibrary_scientist)
    # возвращаем словарь, в котором учёные распределены по группам (ключам словаря)
    return groups


def create_the_third_table():
    doc = Document('Сравнение показателей elibrary.ru and scholar.google.ru.docx')
    doc.add_paragraph('Получив список учёных, у которых есть профили на обоих сайтах, (таблица 2) пройдёмся отдельно '
                      'по профилю каждого учёного на двух сайтах, получим по 5 его самых цитируемых публикаций на '
                      'elibrary.ru и на scholar.google.ru, сравним их на совпадения между друг другом и распределим '
                      'учёных на основе количества совпадений по группам (Если у учёного менее 5 публикаций на одном '
                      'из сайтов, то в 3 таблицу он не попадает).')
    doc.add_paragraph('Таблица 3. Распределение учёных (из таблицы 2) по группам (столбцам):')

    table = doc.add_table(rows=1, cols=4)

    # применяем стиль для таблицы
    table.style = 'Table Grid'

    second_column = table.cell(0, 0)
    second_column.text = 'Самые цитируемые публикации полностью совпадают'

    third_column = table.cell(0, 1)
    third_column.text = 'Самые цитируемые публикации почти совпадают'

    fourth_column = table.cell(0, 2)
    fourth_column.text = 'Самые цитируемые публикации почти не совпадают'

    fifth_column = table.cell(0, 3)
    fifth_column.text = 'Самые цитируемые публикации полностью не совпадают'

    doc.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')

    return doc, table


def fill_in_the_third_table_with_information_from_groups(doc, table, groups):
    # добавляем строку в таблицу (в конец) ровно столько раз, сколько учёных в самой большой группе
    for i in range(1, len(groups[max(groups)]) + 1):
        table.add_row()
    for i in range(1, len(groups['Самые цитируемые публикации полностью совпадают']) + 1):
        # записываем в 1 столбец учёных из группы "5/5"
        cell = table.cell(i, 0)
        cell.text = groups['Самые цитируемые публикации полностью совпадают'][i - 1]
    for i in range(1, len(groups['Самые цитируемые публикации почти совпадают']) + 1):
        # записываем во 2 столбец учёных из групп "4/5" или "3/5"
        cell = table.cell(i, 1)
        cell.text = groups['Самые цитируемые публикации почти совпадают'][i - 1]
    for i in range(1, len(groups['Самые цитируемые публикации почти не совпадают']) + 1):
        # записываем в 3 столбец учёных из групп "1/5" или "2/5"
        cell = table.cell(i, 2)
        cell.text = groups['Самые цитируемые публикации почти не совпадают'][i - 1]
    for i in range(1, len(groups['Самые цитируемые публикации полностью не совпадают']) + 1):
        # записываем во 4 столбец учёных из группы "0/5"
        cell = table.cell(i, 3)
        cell.text = groups['Самые цитируемые публикации полностью не совпадают'][i - 1]
    doc.save('Сравнение показателей elibrary.ru and scholar.google.ru.docx')


def process_second_table_and_create_third_table_based_on_it():
    # получаем словарь, где ключ - учёный, значение ключа - 5 его самых цитируемых публикаций на elibrary.ru
    elibrary_dictionary = get_more_extensive_elibrary_scientists_information()

    # получаем словарь, где ключ - учёный, значение ключа - 5 его самых цитируемых публикаций на scholar.google.ru
    scholar_dictionary = get_more_extensive_scholar_scientists_information()

    # сравнив elibrary_dictionary и scholar_dictionary, получаем словарь, в котором по группам распределены учёные
    groups = compare_dictionaries(elibrary_dictionary, scholar_dictionary)

    # создаём 3-ую таблицу, в которой каждый столбец - определённая группа учёных
    doc, table = create_the_third_table()

    # заполняем столбцы таблицы учёными на основе распределения их по группам
    fill_in_the_third_table_with_information_from_groups(doc, table, groups)


def main():
    # получаем общую информацию о всех учёных на elibrary.ru
    get_information_from_elibrary()
    sleep(3)

    # получаем общую информацию о всех учёных на scholar.google.ru
    get_information_from_google_scholar()
    sleep(3)

    # заполняем выше полученной информацией первую и вторую таблицы (1 табл. - все учёные, 2 табл. - только те,
    # чьи профили есть на обоих сайтах
    process_first_table_and_create_second_table_based_on_it()
    sleep(3)

    # проходимся по тем учёным, чьи профили есть на обоих сайтах: получаем по 5 самых цитируемых статей учёного с
    # обоих сайтов и сравниваем статьи на совпадения, после чего распределяем учёных по группам на основе совпадений
    # их статей
    process_second_table_and_create_third_table_based_on_it()


if __name__ == "__main__":
    main()
